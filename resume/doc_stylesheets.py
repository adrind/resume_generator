import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

EMU_PER_INCH = 91440

def format_doc_1_address_line(resume):
    address = resume('Address')
    city = resume('City')
    email = resume('Email')
    phone = resume('Phone')

    line = address

    if line and city:
        line = line + ' ' + city
    else:
        line = city

    if line and email:
        line = line + ' | ' + email
    else:
        line = email

    if line and phone:
        line = line + ' | ' + phone
    else:
        line = phone

    return line

def set_doc_1_styles(styles):
    styles['Heading 1'].font.size = Pt(28)
    styles['Heading 1'].font.color.rgb = RGBColor(0x2E, 0x7B, 0x86)
    styles['Heading 1'].font.underline = True
    styles['Heading 1'].font.bold = False

    styles.add_style('Address Line', WD_STYLE_TYPE.PARAGRAPH)
    styles['Address Line'].font.size = Pt(12)
    styles['Address Line'].font.bold = True

    styles.add_style('Objective', WD_STYLE_TYPE.PARAGRAPH)
    styles['Objective'].font.italic = True
    styles['Objective'].font.size = Pt(12)

    styles.add_style('Section Header', WD_STYLE_TYPE.PARAGRAPH)
    styles['Section Header'].font.size = Pt(18)
    styles['Section Header'].font.color.rgb = RGBColor(0x2E, 0x7B, 0x86)
    styles['Section Header'].font.underline = True

    styles.add_style('Section Subheader', WD_STYLE_TYPE.PARAGRAPH)
    styles['Section Subheader'].font.size = Pt(12)
    styles['Section Subheader'].font.bold = True
    styles['Section Subheader'].font.all_caps = True


def set_doc_2_styles(styles):
    styles['Heading 1'].font.size = Pt(16)
    first_col_style = styles.add_style('First Col', WD_STYLE_TYPE.PARAGRAPH)
    first_col_style.base_style = styles['Normal']
    first_col_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def init_resume_data(resume):
    return lambda id: next(field for field in resume if field['id'] == id)['data']

def fetch_section(resume, field_id):
    return next(field for field in resume if field['id'] == field_id)['data']

def build_doc_1(resume):
    fetch_field = init_resume_data(resume)

    file_name = os.path.dirname(os.path.abspath(__file__))+'/tmp/test.docx'
    document = Document()

    set_doc_1_styles(document.styles)

    document.add_paragraph(fetch_field('Name'), style='Heading 1')

    document.add_paragraph(format_doc_1_address_line(fetch_field), style='Address Line')
    objective = document.add_paragraph(fetch_field('Objective')['values'][0]['data'], style='Objective')

    objective.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph('Education & Certificates', style='Section Header')

    for ed in fetch_section(resume, 'Education')['values']:
        document.add_paragraph(ed['name']['data'] + ' | ' + ed['dates']['data'], style='Section Subheader')
        for item in ed['description']['data']:
            document.add_paragraph(item, style='ListBullet')

    document.add_paragraph('Experience', 'Section Header')

    for work in fetch_section(resume, 'Work')['values']:
        document.add_paragraph(work['name']['data'] + ' | ' + work['title']['data'] + ' | ' + work['dates']['data'], style='Section Subheader')
        for item in work['description']['data']:
            document.add_paragraph(item, style='ListBullet')

    document.add_paragraph('Skills', style='Section Header')
    for skill in fetch_section(resume, 'Skills')['values'][0]['data']:
        document.add_paragraph(skill, style='ListBullet')

    document.save(file_name)

    return file_name

def build_doc_2(resume):
    fetch_field = init_resume_data(resume)
    file_name = os.path.dirname(os.path.abspath(__file__))+'/tmp/test.docx'
    document = Document()

    document.add_heading(fetch_field('Name'), 0)
    table = document.add_table(rows=2, cols=2)

    address, email = table.rows[0].cells
    city, phone = table.rows[1].cells

    address.text = fetch_field('Address')
    email.paragraphs[0].text = fetch_field('Email')
    email.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    city.text = fetch_field('City')
    phone.paragraphs[0].text = fetch_field('Phone')
    phone.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.add_heading('SKILLS AND ABILITIES', 1)
    skills_table = document.add_table(rows=1, cols=2)
    first_cell, second_cell = skills_table.rows[0].cells

    for i, skill in enumerate(fetch_section(resume, 'Skills')['values'][0]['data']):
        if not (i % 2):
            first_cell.add_paragraph(skill, style='ListBullet')
        else:
            second_cell.add_paragraph(skill, style='ListBullet')

    document.add_heading('WORK AND EXPERIENCE', 1)

    for work in fetch_section(resume, 'Work')['values']:
        work_table = document.add_table(rows=1, cols=2)
        company, dates = work_table.rows[0].cells

        company.text = work['name']['data'] + ', ' + work['title']['data']
        dates.paragraphs[0].text = work['dates']['data']
        dates.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.add_heading('EDUCATION AND CERTIFICATES', 1)

    for ed in fetch_section(resume, 'Education')['values']:
        education_table = document.add_table(rows=1, cols=2)
        name, dates = education_table.rows[0].cells

        name.text = ed['name']['data']
        dates.paragraphs[0].text = ed['dates']['data']
        dates.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    document.save(file_name)
    return file_name

def build_doc_3(resume):
    fetch_field = init_resume_data(resume)
    file_name = os.path.dirname(os.path.abspath(__file__))+'/tmp/test3.docx'
    document = Document()

    document.add_heading(fetch_field('Name'), 0)
    document.add_heading(fetch_field('Address'), 2)
    document.add_heading(fetch_field('City'), 2)
    document.add_heading(fetch_field('Email'), 2)
    document.add_heading(fetch_field('Phone'), 2)

    table = document.add_table(rows=4, cols=2)
    table.allow_autofit = False
    cols = table.columns
    total_width = cols[0].width + cols[1].width
    first_col_width = EMU_PER_INCH * 2
    second_col_width = total_width - EMU_PER_INCH * 2

    objective, skills, education, work = table.rows

    objective.cells[0].text = 'Objective'
    objective.cells[1].text = fetch_field('Objective')['values'][0]['data']
    objective.cells[0].width = first_col_width
    objective.cells[1].width = second_col_width


    skills.cells[0].text = 'Skills and Abilities'
    skills.cells[0].width = first_col_width
    skills.cells[1].width = second_col_width

    for skill in fetch_section(resume, 'Skills')['values'][0]['data']:
        skills.cells[1].add_paragraph(skill, style='ListBullet')

    education.cells[0].text = 'Education and Certificates'
    education.cells[0].width = first_col_width
    education.cells[1].width = second_col_width

    for ed in fetch_section(resume, 'Education')['values']:
        education.cells[1].add_paragraph(ed['name']['data'])
        education.cells[1].add_paragraph(ed['dates']['data'])
        for item in ed['description']['data']:
            education.cells[1].add_paragraph(item, style='ListBullet')

    work.cells[0].text = 'Work and Experience'
    work.cells[0].width = first_col_width
    work.cells[1].width = second_col_width

    for work_item in fetch_section(resume, 'Work')['values']:
        work.cells[1].add_paragraph(work_item['name']['data'])
        work.cells[1].add_paragraph(work_item['title']['data'])
        work.cells[1].add_paragraph(work_item['dates']['data'])
        for item in work_item['description']['data']:
            work.cells[1].add_paragraph(item, style='ListBullet')

    document.save(file_name)
    return file_name

